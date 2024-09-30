from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from .Total import total_header
from .ConvertFile import convert_file


async def change_file(parameters, id_number):
    doc = Document(f'OfferMaking/Templates/{id_number}.docx')
    table = doc.tables[0]

    for i, param in enumerate(parameters[1:]):
        if i in (0, 1):
            row = table.rows[i + 2].cells
        else:
            row = table.add_row().cells

        row[0].text = str(i + 1)
        row[0].paragraphs[0].runs[0].bold = True

        row[1].text = str(param[0])
        row[2].text = str(param[1])
        row[3].text = param[2].capitalize()
        row[4].text = str(param[3])
        row[5].text = str(param[4])

        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

    await total_header(i, table, parameters)

    file = "OfferMaking/Files/КП.docx"
    doc.save(file)

    return await convert_file()
