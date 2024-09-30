from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


async def sum_with_discount(last_row, parameters, total_sum=None, i1=True):
    if i1:
        total_sum = parameters[1][-1]
    discount = parameters[0]
    final_sum = total_sum - discount

    cell = last_row[-1]
    cell.text = str(final_sum)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


async def total_body(table, parameters, row1):
    total_quantity = sum(int(row.cells[4].text) for row in table.rows[2:-2])
    row1[4].text = str(total_quantity)
    paragraph = row1[4].paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    total_sum = sum(int(row.cells[5].text) for row in table.rows[2:-2])
    row1[5].text = str(total_sum)
    paragraph = row1[5].paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    await sum_with_discount(table.rows[-1].cells, parameters, total_sum, False)


async def total_header(i, table, parameters):
    if i == 0:
        last_row = table.rows[-1].cells
        merged_cell = last_row[0].merge(last_row[4])
        merged_cell.text = "ЦЕНА СО СКИДКОЙ:"

        paragraph = merged_cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)

        await sum_with_discount(last_row, parameters)

    else:
        row1 = table.add_row().cells
        merged_cell1 = row1[0].merge(row1[3])
        merged_cell1.text = "Итого:"

        paragraph1 = merged_cell1.paragraphs[0]
        paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = paragraph1.runs[0]
        run1.bold = True
        run1.font.color.rgb = RGBColor(0, 0, 0)

        row2 = table.add_row().cells
        merged_cell2 = row2[0].merge(row2[4])
        merged_cell2.text = "ИТОГО СО СКИДКОЙ:"

        paragraph2 = merged_cell2.paragraphs[0]
        paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = paragraph2.runs[0]
        run2.bold = True
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(255, 0, 0)

        await total_body(table, parameters, row1)
