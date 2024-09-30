from docx import Document
from docx.shared import Pt, RGBColor, Cm
from Instruments.Config import auth
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import aiohttp
import io


async def company_information(company_data, id_number):
    file_path = f'OfferMaking/Templates/{id_number}.docx'
    if os.path.exists(f'OfferMaking/Templates/{id_number}.docx'):
        doc = Document(file_path)
    else:
        doc = Document('OfferMaking/Files/Шаблон.docx')

    section = doc.sections[0]
    header = section.header
    table = header.tables[0]

    phone, address, email, name = company_data
    text_cell = table.cell(0, 0)

    for paragraph in text_cell.paragraphs:
        if "Номер телефона:" in paragraph.text:
            paragraph.text = f"Номер телефона: {phone}"
            paragraph.runs[0].font.size = Pt(12)
        elif "Адрес:" in paragraph.text:
            paragraph.text = f"Адрес: {address}"
            paragraph.runs[0].font.size = Pt(12)
        elif "Email:" in paragraph.text:
            paragraph.text = f"Email: {email}"
            paragraph.runs[0].font.size = Pt(12)
            break

    for paragraph in doc.paragraphs:
        if "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ от " in paragraph.text:
            paragraph.text = "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ от "
            run = paragraph.add_run(name)
            run.font.color.rgb = RGBColor(0xfc, 0x9a, 0x40)
            run.font.size = Pt(14)
            for run in paragraph.runs:
                run.font.size = Pt(14)
            break

    doc.save(f'OfferMaking/Templates/{id_number}.docx')


async def logo(logo_url, id_number):
    file_path = f'OfferMaking/Templates/{id_number}.docx'
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document('OfferMaking/Files/Шаблон.docx')

    section = doc.sections[0]
    header = section.header

    table = header.tables[0]

    logo_cell = table.cell(0, 1)
    for paragraph in logo_cell.paragraphs:
        logo_cell._element.remove(paragraph._element)
    logo_cell.add_paragraph().add_run().add_picture(await download_file(logo_url), width=Cm(2.6))
    logo_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(f'OfferMaking/Templates/{id_number}.docx')


async def download_file(file_url):
    async with aiohttp.ClientSession() as session:
        async with session.get(file_url, auth=auth) as response:
            file_data = await response.read()
            return io.BytesIO(file_data)
